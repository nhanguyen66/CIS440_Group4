"""
Definition of views.
"""
from django.shortcuts import redirect
from datetime import datetime
from django.shortcuts import render
from django.http import HttpRequest
from app.forms import TextEditorForm
from app.models import TextEditor

from django.http import HttpResponse
from docx import Document
#from docx.shared import Inches

def home(request):
    """Renders the home page."""
    assert isinstance(request, HttpRequest)
    return render(request,
        'app/index.html',
        {
            'title':'Home Page',
            'year':datetime.now().year,
        })

def textbox(request):
    form = TextEditorForm(request.POST or None)

    if request.method == "POST":
       if 'process_text' in request.POST:
            if form.is_valid():
                    #creating new object of model TextEditor
                    obj = TextEditor()
                    #obj.save_as_docx = False

                    #sending form info to obj
                    obj.textinput = form.cleaned_data['textinput']

                    #performing text processing
                    #all functions edit the obj.textoutput variable
                    obj.process_text()
                    #obj.check_spelling() #needs changing
                    obj.check_difficulty()
                    obj.summarize_text()

                    #obj.save_as_docx = True

                    if(False):#'savedoc' in request.POST):#'savedoc' in request.POST):
                        document = Document()
                        document.add_heading('Key Phrases')
                        document.add_page_break()

                        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = 'attachment; filename=download.docx'
                        document.save(response)
                        return response

                    else:
                        #sending to output page with the textoutput from the model
                        return render(request,
                            'app/output.html',
                            {
                                'title':'Output',
                                'message':'Your output page.',
                                'year':datetime.now().year,
                                'output': obj.textoutput,
                                'difficultyReport':obj.difficultyReport,
                                'summarizedText':obj.summarizedText
                            })
    else:
        return render(request, "app/textbox.html", {"form": form})


def contact(request):
    """Renders the contact page."""
    assert isinstance(request, HttpRequest)
    return render(request,
        'app/contact.html',
        {
            'title':'Contact',
            'message':'Your contact page.',
            'year':datetime.now().year,
        })

def about(request):
    """Renders the about page."""
    assert isinstance(request, HttpRequest)
    return render(request,
        'app/about.html',
        {
            'title':'About',
            'message':'Your application description page.',
            'year':datetime.now().year,
        })

def output(request):
    """Renders the output page."""
    assert isinstance(request, HttpRequest)
    return render(request,
        'app/output.html',
        {
            'title':'summarizedText',
            'message':'Your output page.',
            'year':datetime.now().year,
        })

def savedoc(request):
    form = TextEditorForm(request.POST or None)

    if request.method == "POST":
       if 'process_text' in request.POST:
            if form.is_valid():
                    #creating new object of model TextEditor
                    obj = TextEditor()
                    #obj.save_as_docx = False

                    #sending form info to obj
                    obj.textinput = form.cleaned_data['textinput']

                    #performing text processing
                    #all functions edit the obj.textoutput variable
                    obj.process_text()
                    #obj.check_spelling() #needs changing
                    obj.check_difficulty()
                    obj.summarize_text()

                    #obj.save_as_docx = True

                    #if('savedoc' in request.POST):#'savedoc' in request.POST):
                    document = Document()
                    document.add_heading('Key Phrases')
                    for phrase in obj.summarizedText:
                        document.add_paragraph(phrase, style='ListBullet')
                    document.add_heading("Difficulty Report")
                    document.add_paragraph(obj.difficultyReport)
                    document.add_heading("Text - Needs More Spell Check")
                    document.add_paragraph(obj.textoutput)

                    document.add_page_break()

                    if(True):
                        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = 'attachment; filename=download.docx'
                        document.save(response)
                        return response

                    else:
                        #sending to output page with the textoutput from the model
                        return render(request,
                            'app/output.html',
                            {
                                'title':'Output',
                                'message':'Your output page.',
                                'year':datetime.now().year,
                                'output': obj.textoutput,
                                'difficultyReport':obj.difficultyReport,
                                'summarizedText':obj.summarizedText
                            })
    else:
        return render(request, "app/textbox.html", {"form": form})